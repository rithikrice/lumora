"""Simplified document parsers - extract full text for Gemini's long context.

No chunking/embeddings needed - Gemini 1.5/2.0 handles 1M-2M tokens!
"""

from typing import Dict, Any
from pathlib import Path
import io

from ..core.logging import get_logger
from ..core.errors import ProcessingError
from ..core.config import get_settings

logger = get_logger(__name__)


def parse_pdf(content: bytes, filename: str) -> str:
    """Extract full text from PDF.
    
    Args:
        content: PDF bytes
        filename: Original filename
        
    Returns:
        Full extracted text
    """
    settings = get_settings()
    text = ""
    
    try:
        if settings.USE_VERTEX:
            # Use Document AI for best accuracy
            text = _parse_pdf_with_document_ai(content)
        else:
            # Use PyMuPDF locally
            text = _parse_pdf_with_pymupdf(content)
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        # Try fallback
        try:
            text = _parse_pdf_with_pdfplumber(content)
        except Exception as e2:
            logger.error(f"PDF fallback failed: {e2}")
            text = f"[PDF parsing failed for {filename}]"
    
    return text


def _parse_pdf_with_pymupdf(content: bytes) -> str:
    """Extract text using PyMuPDF."""
    import fitz  # PyMuPDF
    
    text_parts = []
    pdf = fitz.open(stream=content, filetype="pdf")
    
    for page_num, page in enumerate(pdf, 1):
        page_text = page.get_text()
        
        if page_text.strip():
            text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")
    
    pdf.close()
    return "\n".join(text_parts)


def _parse_pdf_with_pdfplumber(content: bytes) -> str:
    """Extract text using pdfplumber."""
    import pdfplumber
    
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")
    
    return "\n".join(text_parts)


def _parse_pdf_with_document_ai(content: bytes) -> str:
    """Extract text using Google Document AI."""
    from google.cloud import documentai_v1 as documentai
    
    settings = get_settings()
    client = documentai.DocumentProcessorServiceClient()
    
    # Use default OCR processor
    name = f"projects/{settings.GOOGLE_PROJECT_ID}/locations/{settings.GOOGLE_LOCATION}/processors/ocr-processor"
    
    request = documentai.ProcessRequest(
        name=name,
        raw_document=documentai.RawDocument(
            content=content,
            mime_type="application/pdf"
        )
    )
    
    result = client.process_document(request=request)
    return result.document.text


def parse_text(content: bytes, filename: str) -> str:
    """Extract text from text file.
    
    Args:
        content: Text bytes
        filename: Original filename
        
    Returns:
        Full text
    """
    try:
        return content.decode('utf-8', errors='replace')
    except Exception as e:
        logger.error(f"Text parsing failed: {e}")
        return f"[Text parsing failed for {filename}]"


def parse_docx(content: bytes, filename: str) -> str:
    """Extract text from DOCX.
    
    Args:
        content: DOCX bytes
        filename: Original filename
        
    Returns:
        Full text
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return f"[DOCX parsing failed for {filename}]"


def parse_document(content: bytes, filename: str) -> str:
    """Parse document and extract full text.
    
    No chunking - returns complete document text for Gemini.
    
    Args:
        content: Document bytes
        filename: Original filename
        
    Returns:
        Full document text
    """
    ext = Path(filename).suffix.lower()
    
    if ext == '.pdf':
        return parse_pdf(content, filename)
    elif ext in ['.txt', '.md']:
        return parse_text(content, filename)
    elif ext == '.docx':
        return parse_docx(content, filename)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return f"[Unsupported file type: {filename}]"


def load_startup_context(startup_id: str, db_service) -> str:
    """Load full startup context for Gemini.
    
    Combines questionnaire responses + all uploaded documents.
    
    Args:
        startup_id: Startup ID
        db_service: Database service
        
    Returns:
        Full context string
    """
    context_parts = []
    
    # Get questionnaire data
    startup_data = db_service.get_startup(startup_id)
    
    if startup_data:
        responses = startup_data.get("questionnaire_responses", {})
        
        context_parts.append("=== QUESTIONNAIRE RESPONSES ===\n")
        
        for key, value in responses.items():
            if value:
                # Format key nicely
                label = key.replace('_', ' ').title()
                context_parts.append(f"{label}: {value}")
        
        context_parts.append("\n")
    
    # Get uploaded documents (from GCS or local)
    # TODO: Implement document retrieval from storage
    # For now, questionnaire is enough for most analyses
    
    return "\n".join(context_parts)

