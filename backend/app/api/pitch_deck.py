"""Enhanced Pitch Deck Upload with Gemini 2.5 + Document AI.

Extracts text, images, graphs, charts, and diagrams from pitch decks.
"""

import base64
import io
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form
from pydantic import BaseModel
import google.generativeai as genai

from ..core.security import verify_api_key
from ..core.logging import get_logger
from ..core.config import get_settings
from ..services.database import get_database_service
from ..services.gcs import GCSService

router = APIRouter()
logger = get_logger(__name__)
settings = get_settings()


class PitchDeckUploadResponse(BaseModel):
    """Response from pitch deck upload."""
    success: bool
    startup_id: str
    document_id: str
    pages_processed: int
    text_extracted: bool
    images_extracted: int
    gemini_analysis: Optional[str] = None


class ExtractedPage(BaseModel):
    """Extracted data from a single page."""
    page_number: int
    text: str
    has_chart: bool
    has_diagram: bool
    visual_description: Optional[str] = None
    key_metrics: List[str] = []


@router.post("/pitch-deck/upload")
async def upload_pitch_deck(
    file: UploadFile = File(...),
    company_name: Optional[str] = Form(None),
    extract_visuals: bool = Form(True),
    doc_type: str = Form("pitch_deck"),
    api_key: str = Depends(verify_api_key)
) -> PitchDeckUploadResponse:
    """Upload and analyze documents with Gemini 2.5.
    
    Supports: PDF, DOCX
    
    Extracts:
    - All text content
    - Charts and graphs (described by Gemini)
    - Financial metrics
    - Key insights
    
    Args:
        file: Document file (PDF or DOCX)
        startup_id: Startup identifier (auto-generated if not provided)
        company_name: Optional company name (extracted from filename if not provided)
        extract_visuals: Whether to extract and analyze visual elements
        doc_type: Type of document (pitch_deck, checklist, investment_memo, financial_report)
        api_key: API key
        
    Returns:
        Upload response with analysis
    """
    import re
    # Auto-generate startup_id if not provided
    if not company_name:
        name_from_file = (file.filename or "").rsplit(".", 1)[0]
        name_from_file = re.sub(r'[-_](pitch|deck|presentation|final|v\d+).*', '', name_from_file, flags=re.IGNORECASE)
        company_name = name_from_file.strip() or None

    # Ask the DB (Firestore if enabled) if this company already exists
    db = get_database_service()
    existing_id = None
    if hasattr(db, "find_startup_id"):
        existing_id = db.find_startup_id(company_name or "")

    if existing_id:
        startup_id = existing_id
    else:
        # deterministically slugify; only fall back if no name
        if company_name:
            safe = re.sub(r'[^a-z0-9]+', '-', company_name.lower()).strip('-')
            startup_id = safe or "startup"
        else:
            # last resort when we have no reasonable name
            import time
            startup_id = f"startup-{int(time.time())}"

    logger.info(f"Resolved startup_id: {startup_id}  company_name: {company_name}")

    
    logger.info(f"Uploading {doc_type} for {startup_id}: {file.filename}")
    
    try:
        # Validate file type
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.pdf') or filename_lower.endswith('.docx')):
            raise HTTPException(400, "Only PDF and DOCX files are supported")
        
        # Read content
        content = await file.read()
        size_mb = len(content) / (1024 * 1024)
        
        if size_mb > 50:  # 50MB limit
            raise HTTPException(400, "Pitch deck must be under 50MB")
        
        logger.info(f"Processing {size_mb:.1f}MB {file.filename}")
        
        # Extract text and images based on file type
        if filename_lower.endswith('.pdf'):
            pages_data = await _extract_pdf_with_gemini(
                content=content,
                extract_visuals=extract_visuals
            )
        elif filename_lower.endswith('.docx'):
            pages_data = await _extract_docx_with_gemini(
                content=content,
                extract_visuals=extract_visuals
            )
        else:
            raise HTTPException(400, "Unsupported file type")
        
        # Generate document ID
        import hashlib
        document_id = f"pitch-{startup_id}-{hashlib.md5(content).hexdigest()[:8]}"
        
        # Upload to GCS
        gcs_service = GCSService()
        storage_path, public_url = await gcs_service.upload_file(
            content=content,
            filename=file.filename,
            startup_id=startup_id
        )
        
        # Store extracted data in database
        db = get_database_service()
        startup_data = db.get_startup(startup_id)
        
        # Store document data
        document_data = {
            "document_id": document_id,
            "filename": file.filename,
            "doc_type": doc_type,
            "storage_path": storage_path,
            "public_url": public_url,
            "pages": [page.dict() for page in pages_data],
            "full_text": "\n\n".join([p.text for p in pages_data]),
            "total_pages": len(pages_data),
            "images_count": sum(1 for p in pages_data if p.has_chart or p.has_diagram)
        }
        
        # Store in appropriate field based on doc_type
        doc_field = doc_type if doc_type != "pitch_deck" else "pitch_deck"
        
        # Always merge with existing data to avoid overwriting
        existing_responses = {}
        if startup_data:
            existing_responses = startup_data.get("questionnaire_responses", {})
        
        # Set company_name: prioritize provided name > existing name > extracted from filename
        if company_name:
            existing_responses["company_name"] = company_name
        elif "company_name" not in existing_responses or not existing_responses["company_name"]:
            # Extract from filename as fallback
            name_from_file = file.filename.rsplit('.', 1)[0]
            import re
            name_from_file = re.sub(r'[-_](pitch|deck|presentation|final|v\d+).*', '', name_from_file, flags=re.IGNORECASE)
            existing_responses["company_name"] = name_from_file.strip() or startup_id
        
        # Add/update the pitch deck data
        existing_responses[doc_field] = document_data
        
        # Save merged data
        db.save_questionnaire_response(
            startup_id=startup_id,
            responses=existing_responses
        )
        
        logger.info(f"Merged pitch deck data into startup {startup_id}, existing fields: {list(existing_responses.keys())}")
        
        # Extract structured profile using unified extractor (LLM + regex fallback)
        from ..services.extractors import extract_structured_profile
        structured = await extract_structured_profile(
            pages=[p.dict() for p in pages_data],
            api_key=settings.GEMINI_API_KEY,
            pdf_bytes=content,
            doc_type="pitch_deck"
        )
        
        # Set company name if provided and not already in structured data
        if company_name and not structured.get("company_name"):
            structured["company_name"] = company_name
        
        # Write-through to profile (SSoT)
        profile = db.save_structured_profile(startup_id, structured, source="pitch_deck")
        
        # Keep an index of documents (lightweight)
        doc_index = {
            "document_id": document_id,
            "type": doc_type,
            "filename": file.filename,
            "public_url": public_url,
            "pages": len(pages_data)
        }
        db.add_document_index(startup_id, doc_index)
        
        # Generate quick analysis with Gemini 2.5
        analysis = await _analyze_pitch_deck_with_gemini(pages_data)
        
        logger.info(f"Pitch deck processed: {document_id}, {len(pages_data)} pages")
        
        return PitchDeckUploadResponse(
            success=True,
            startup_id=startup_id,
            document_id=document_id,
            pages_processed=len(pages_data),
            text_extracted=True,
            images_extracted=sum(1 for p in pages_data if p.has_chart or p.has_diagram),
            gemini_analysis=analysis
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Pitch deck upload error: {e}", exc_info=True)
        raise HTTPException(500, f"Upload failed: {str(e)}")


async def _extract_pdf_with_gemini(
    content: bytes,
    extract_visuals: bool = True
) -> List[ExtractedPage]:
    """Extract text and visuals from PDF using Gemini 2.5.
    
    Args:
        content: PDF bytes
        extract_visuals: Whether to analyze visual elements
        
    Returns:
        List of extracted pages
    """
    import fitz  # PyMuPDF
    
    pages_data = []
    pdf = fitz.open(stream=content, filetype="pdf")
    
    logger.info(f"Processing PDF with {len(pdf)} pages")
    
    for page_num, page in enumerate(pdf, 1):
        # Extract text
        text = page.get_text()
        
        has_chart = False
        has_diagram = False
        visual_description = None
        key_metrics = []
        
        # If extract_visuals is enabled, analyze page image with Gemini
        if extract_visuals and settings.GEMINI_API_KEY:
            try:
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution
                img_bytes = pix.tobytes("png")
                
                # Analyze with Gemini 2.5 Flash
                visual_analysis = await _analyze_page_with_gemini(
                    img_bytes=img_bytes,
                    page_text=text,
                    page_num=page_num
                )
                
                has_chart = visual_analysis.get("has_chart", False)
                has_diagram = visual_analysis.get("has_diagram", False)
                visual_description = visual_analysis.get("description")
                key_metrics = visual_analysis.get("metrics", [])
                
            except Exception as e:
                logger.warning(f"Visual analysis failed for page {page_num}: {e}")
        
        pages_data.append(ExtractedPage(
            page_number=page_num,
            text=text,
            has_chart=has_chart,
            has_diagram=has_diagram,
            visual_description=visual_description,
            key_metrics=key_metrics
        ))
    
    pdf.close()
    return pages_data


async def _extract_docx_with_gemini(
    content: bytes,
    extract_visuals: bool = True
) -> List[ExtractedPage]:
    """Extract text from DOCX using python-docx.
    
    Args:
        content: DOCX bytes
        extract_visuals: Whether to analyze visual elements (images in DOCX)
        
    Returns:
        List of extracted pages (treated as sections)
    """
    from docx import Document
    import io
    
    pages_data = []
    doc = Document(io.BytesIO(content))
    
    logger.info(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")
    
    # Extract all text
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    # Split into logical sections (by headings or every ~500 words)
    sections = []
    current_section = []
    word_count = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        # Check if it's a heading
        is_heading = para.style.name.startswith('Heading') if para.style else False
        
        if is_heading and current_section:
            sections.append("\n".join(current_section))
            current_section = []
            word_count = 0
        
        current_section.append(para.text)
        word_count += len(para.text.split())
        
        # Split every ~500 words
        if word_count > 500:
            sections.append("\n".join(current_section))
            current_section = []
            word_count = 0
    
    if current_section:
        sections.append("\n".join(current_section))
    
    # Create ExtractedPage for each section
    for idx, section_text in enumerate(sections, 1):
        # Analyze with Gemini if extract_visuals is enabled
        key_metrics = []
        if extract_visuals and settings.GEMINI_API_KEY:
            try:
                key_metrics = await _extract_metrics_from_text(section_text)
            except Exception as e:
                logger.warning(f"Metric extraction failed for section {idx}: {e}")
        
        pages_data.append(ExtractedPage(
            page_number=idx,
            text=section_text,
            has_chart=False,  # DOCX text extraction doesn't detect charts
            has_diagram=False,
            visual_description=None,
            key_metrics=key_metrics
        ))
    
    return pages_data


async def _extract_metrics_from_text(text: str) -> List[str]:
    """Extract financial metrics from text using Gemini.
    
    Args:
        text: Text to analyze
        
    Returns:
        List of extracted metrics
    """
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-pro")  # Using Gemini 2.5 Pro for better metric extraction
        
        prompt = f"""Extract financial and business metrics from this text.
Return only the metrics in a list format.

Text:
{text[:1000]}

Examples of metrics to extract:
- ARR: $5M
- Growth: 200%
- Customers: 500
- Team size: 25

Return just the list, one per line."""
        
        response = model.generate_content(prompt)
        metrics_text = response.text.strip()
        
        # Parse metrics (each line is a metric)
        metrics = [line.strip() for line in metrics_text.split('\n') if line.strip() and not line.strip().startswith('-')]
        return metrics[:10]  # Limit to 10 metrics
        
    except Exception as e:
        logger.error(f"Metric extraction failed: {e}")
        return []


async def _analyze_page_with_gemini(
    img_bytes: bytes,
    page_text: str,
    page_num: int
) -> Dict[str, Any]:
    """Analyze a single page image with Gemini 2.5 Flash.
    
    Args:
        img_bytes: Page image as PNG
        page_text: Extracted text
        page_num: Page number
        
    Returns:
        Analysis results
    """
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-pro")  # Using Gemini 2.5 Pro with vision for better chart analysis
        
        # Prepare image
        import PIL.Image
        img = PIL.Image.open(io.BytesIO(img_bytes))
        
        prompt = f"""Analyze this pitch deck slide (page {page_num}).

Text content: {page_text[:500]}

Identify:
1. Does it contain charts or graphs? (revenue, growth, metrics)
2. Does it contain diagrams or flowcharts?
3. Extract any numerical metrics (ARR, MRR, growth %, etc.)
4. Provide a 1-sentence description of visual elements

Return JSON:
{{
  "has_chart": true/false,
  "has_diagram": true/false,
  "metrics": ["ARR: $5M", "Growth: 200%"],
  "description": "Revenue growth chart showing 3x increase"
}}
"""
        
        response = model.generate_content([prompt, img])
        
        # Parse JSON response
        import json
        import re
        
        response_text = response.text
        # Extract JSON from markdown if needed
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            result = json.loads(json_match.group())
        else:
            result = {
                "has_chart": False,
                "has_diagram": False,
                "metrics": [],
                "description": response_text[:200]
            }
        
        return result
        
    except Exception as e:
        logger.error(f"Gemini page analysis failed: {e}")
        return {
            "has_chart": False,
            "has_diagram": False,
            "metrics": [],
            "description": None
        }


async def _analyze_pitch_deck_with_gemini(pages_data: List[ExtractedPage]) -> str:
    """Generate quick pitch deck analysis with Gemini.
    
    Args:
        pages_data: Extracted pages
        
    Returns:
        Summary analysis
    """
    if not settings.GEMINI_API_KEY:
        return "Analysis skipped: Gemini API key not configured"
    
    # Check if API key looks like OpenAI key (starts with sk-)
    if settings.GEMINI_API_KEY.startswith('sk-'):
        logger.error(f"Invalid Gemini API key format (looks like OpenAI key)")
        return "Analysis skipped: Invalid API key format (use Gemini API key, not OpenAI)"
    
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        
        # Use Gemini 2.5 Pro for structured extraction
        model = genai.GenerativeModel("gemini-2.5-pro")
        
        # Combine all text (limit to avoid token limits)
        full_text = "\n\n".join([
            f"Page {p.page_number}: {p.text[:500]}"  # Limit per page
            for p in pages_data[:10]  # Only first 10 pages
        ])
        
        prompt = f"""Analyze this pitch deck in 3 sentences:
1. What the company does
2. Key metrics/traction
3. Value proposition

Content:
{full_text[:4000]}"""
        
        # Set timeout and limit tokens
        import asyncio
        response = await asyncio.wait_for(
            asyncio.to_thread(model.generate_content, prompt),
            timeout=20.0
        )
        
        return response.text.strip()
        
    except asyncio.TimeoutError:
        logger.warning("Pitch deck analysis timed out")
        return "Analysis skipped: Timeout"
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Pitch deck analysis failed: {error_msg}", exc_info=True)
        
        # Check for common errors
        if "API key" in error_msg or "authentication" in error_msg.lower():
            return "Analysis skipped: API key authentication failed"
        elif "quota" in error_msg.lower():
            return "Analysis skipped: API quota exceeded"
        else:
            return f"Analysis skipped: {error_msg[:100]}"


@router.get("/pitch-deck/{startup_id}")
async def get_pitch_deck_data(
    startup_id: str,
    api_key: str = Depends(verify_api_key)
):
    """Get extracted pitch deck data.
    
    Args:
        startup_id: Startup identifier
        api_key: API key
        
    Returns:
        Pitch deck data
    """
    db = get_database_service()
    startup_data = db.get_startup(startup_id)
    
    if not startup_data:
        raise HTTPException(404, "Startup not found")
    
    # Check questionnaire_responses for pitch deck
    responses = startup_data.get("questionnaire_responses", {})
    if "pitch_deck" not in responses:
        raise HTTPException(404, "Pitch deck not found for this startup")
    
    return {
        "startup_id": startup_id,
        "pitch_deck": responses["pitch_deck"]
    }

